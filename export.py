import io
import html as _html
import markdown as _markdown

SECTION_LABELS = {
    "market": "Market Context",
    "performance": "Portfolio Performance",
    "risk": "Risk Analysis",
}

def _period_of(result: dict) -> str:
    plan = result.get("plan") or {}
    return (
        plan.get("period")
        or (result.get("performance") or {}).get("period")
        or ""
    )

def _title_of(result: dict) -> str:
    plan = result.get("plan") or {}
    intent = (plan.get("intent") or result.get("response_type") or "report").replace("_", " ").title()
    period = _period_of(result)
    return f"Portfolio AI — {intent}" + (f": {period}" if period else "")

def _meta_line(result: dict) -> str | None:
    summary = (result.get("performance") or {}).get("pnl_summary")
    if not summary or summary.get("return_pct") is None:
        return None
    ret = summary["return_pct"]
    start = summary.get("start_aum", 0)
    end = summary.get("end_aum", 0)
    total = summary.get("total_pnl", 0)
    return (
        f"Portfolio Return: {ret:+.2f}%  |  Start AUM: ${start:,.0f}  |  "
        f"End AUM: ${end:,.0f}  |  Total P&L: ${total:,.0f}"
    )

def _assemble_blocks(result: dict) -> list[tuple[str, str]]:
    blocks: list[tuple[str, str]] = [("title", _title_of(result))]

    meta = _meta_line(result)
    if meta:
        blocks.append(("meta", meta))

    # newsletter
    if result.get("newsletter", {}).get("newsletter"):
        blocks.append(("heading", "Newsletter"))
        blocks.append(("md", result["newsletter"]["newsletter"]))

    # sections
    for key in ("market", "performance", "risk"):
        section = result.get(key)
        if not section or not section.get("analysis"):
            continue
        blocks.append(("heading", SECTION_LABELS[key]))
        # Risk carries a pre-aligned metrics block — show it as monospace first
        if key == "risk" and section.get("metrics"):
            blocks.append(("mono", section["metrics"]))
        blocks.append(("md", section["analysis"]))

    # stress test
    st_result = result.get("stress_test")
    if st_result and not st_result.get("error"):
        scen = ", ".join(st_result.get("scenarios_used", []))
        blocks.append(("heading", f"Stress Test — {scen}" if scen else "Stress Test"))
        if st_result.get("tables"):
            blocks.append(("mono", st_result["tables"]))
        if st_result.get("narrative"):
            blocks.append(("md", st_result["narrative"]))

    return blocks

# ----------------------------------------------------------
# PDF renderer — blocks → HTML → xhtml2pdf
# ----------------------------------------------------------

PDF_CSS = """
@page { size: letter; margin: 2cm; }
body { font-family: "Times New Roman", Times, serif; font-size: 11pt; color: #000; line-height: 1.4; }
h1 { font-size: 18pt; margin: 0 0 4pt 0; }
h2 { font-size: 14pt; margin: 16pt 0 6pt 0; border-bottom: 1px solid #000; padding-bottom: 2pt; }
h3 { font-size: 12pt; margin: 10pt 0 4pt 0; }
.meta { font-size: 10pt; color: #222; margin: 0 0 10pt 0; }
table { border-collapse: collapse; width: 100%; margin: 6pt 0; }
th, td { border: 1px solid #444; padding: 3pt 5pt; font-size: 9.5pt; text-align: left; }
pre { font-family: "Courier New", monospace; font-size: 8pt; white-space: pre; margin: 6pt 0; }
"""

def _blocks_to_html(blocks: list[tuple[str, str]]) -> str:
    parts = ['<html><head><meta charset="utf-8">', f"<style>{PDF_CSS}</style></head><body>"]
    for kind, text in blocks:
        if kind == "title":
            parts.append(f"<h1>{_html.escape(text)}</h1>")
        elif kind == "meta":
            parts.append(f'<div class="meta">{_html.escape(text)}</div>')
        elif kind == "heading":
            parts.append(f"<h2>{_html.escape(text)}</h2>")
        elif kind == "mono":
            parts.append(f"<pre>{_html.escape(text)}</pre>")
        elif kind == "md":
            parts.append(_markdown.markdown(text, extensions=["tables", "sane_lists"]))
    parts.append("</body></html>")
    return "\n".join(parts)


def build_pdf(result: dict) -> bytes:
    from xhtml2pdf import pisa
    html_doc = _blocks_to_html(_assemble_blocks(result))
    buf = io.BytesIO()
    pisa.CreatePDF(src=html_doc, dest=buf, encoding="utf-8")
    return buf.getvalue()


# ----------------------------------------------------------
# DOCX renderer — blocks → python-docx/htmldocx
# ----------------------------------------------------------

def build_docx(result: dict) -> bytes:
    from docx import Document
    from docx.shared import Pt
    from htmldocx import HtmlToDocx

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)

    parser = HtmlToDocx()
    blocks = _assemble_blocks(result)

    for kind, text in blocks:
        if kind == "title":
            doc.add_heading(text, level=0)
        elif kind == "meta":
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.italic = True
            run.font.size = Pt(10)
        elif kind == "heading":
            doc.add_heading(text, level=1)
        elif kind == "mono":
            _add_monospace(doc, text, Pt)
        elif kind == "md":
            html_frag = _markdown.markdown(text, extensions=["tables", "sane_lists"])
            parser.add_html_to_document(html_frag, doc)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

def _add_monospace(doc, text: str, Pt):
    p = doc.add_paragraph()
    lines = text.split("\n")
    for i, line in enumerate(lines):
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(8)
        if i < len(lines) - 1:
            run.add_break()

def export_filename(result: dict, ext: str) -> str:
    plan = result.get("plan") or {}
    rtype = (plan.get("intent") or result.get("response_type") or "report")
    period = _period_of(result) or "report"
    return f"portfolio_{period}_{rtype}.{ext}".replace(" ", "_")
